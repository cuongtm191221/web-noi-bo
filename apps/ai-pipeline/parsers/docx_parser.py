import io
from .base import ParsedDocument, PageContent


def parse_docx(file_bytes: bytes) -> ParsedDocument:
    """Parse DOCX with python-docx."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return ParsedDocument(
        pages=[PageContent(text=full_text, location={"page_number": 1})],
        metadata={"format": "docx", "paragraph_count": len(paragraphs)},
    )